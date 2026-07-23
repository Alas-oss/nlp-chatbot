from pathlib import Path
from docx import Document as DocxDocument
from langchain_core.documents import Document
from langchain_text_splitters.character import RecursiveCharacterTextSplitter
from pdf_loader import load_pdf_as_documents


def load_docx_as_documents(path: str) -> list[Document]:
    docx = DocxDocument(path)
    full_text = "\n\n".join(p.text for p in docx.paragraphs if p.text.strip())
    return [Document(page_content=full_text, metadata={"source": path})]


def load_all_sources(folder_path: str) -> list[Document]:
    all_docs = []
    for path in Path(folder_path).glob("*.docx"):
        docs = load_docx_as_documents(str(path))
        for d in docs:
            d.metadata["doc_type"] = "docx"
        all_docs.extend(docs)
    for path in Path(folder_path).glob("*.pdf"):
        docs = load_pdf_as_documents(str(path))
        for d in docs:
            d.metadata["doc_type"] = "pdf"
        all_docs.extend(docs)
    return all_docs


def chunk_documents(documents: list[Document]) -> list[Document]:
    splitter = RecursiveCharacterTextSplitter(
        chunk_size=600, chunk_overlap=80,
        separators=["\n\n", "\n", ". ", " ", ""],
    )
    return splitter.split_documents(documents)