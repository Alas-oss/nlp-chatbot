from docx import Document

doc = Document()
doc.add_heading("King's College London - General Information", level=0)

doc.add_heading("History and Founding", level=1)
doc.add_paragraph(
    "King's College London (KCL) was founded in 1829 by King George IV and the "
    "Duke of Wellington, making it one of the oldest universities in England. It "
    "was established, in part, as a Church of England response to the founding of "
    "the secular University College London a few years earlier. King's became one "
    "of the founding constituent colleges of the University of London in 1836. "
    "The university's motto is 'Sancte et Sapienter' — 'With Holiness and Wisdom'."
)

doc.add_heading("Campuses", level=1)
doc.add_paragraph(
    "King's operates across five main campuses in central London: the Strand Campus "
    "(the historic main site, home to arts, humanities, artificial intelligence, and " \
    "law), Guy's Campus (near London Bridge, focused on health and life sciences), St " \
    "Thomas' Campus (opposite the Houses of Parliament, home to nursing and midwifery), " \
    "Waterloo Campus (social sciences, public policy, computer science, and natural " \
    "sciences), and Denmark Hill Campus (south London, focused on medical, dental, and " \
    "psychiatric research alongside King's College Hospital)."
)

doc.add_heading("Faculties and Schools", level=1)
doc.add_paragraph(
    "King's is organised into nine faculties: the Faculty of Arts & Humanities; "
    "the Faculty of Dentistry, Oral & Craniofacial Sciences; the Faculty of Life "
    "Sciences & Medicine; the Faculty of Natural, Mathematical & Engineering "
    "Sciences; the Faculty of Nursing, Midwifery & Palliative Care; the Faculty "
    "of Social Science & Public Policy; the Institute of Psychiatry, Psychology "
    "& Neuroscience; King's Business School; and The Dickson Poon School of Law."
)

doc.add_heading("Notable History and Contributions", level=1)
doc.add_paragraph(
    "King's has a strong historical connection to major scientific and medical "
    "milestones. Florence Nightingale founded the world's first official nursing "
    "school at King's in 1860. Researchers at King's, including Rosalind "
    "Franklin and Maurice Wilkins, produced X-ray diffraction images central to "
    "discovering the double-helix structure of DNA in the 1950s. The Nightingale "
    "legacy continues through the Florence Nightingale Faculty of Nursing, "
    "Midwifery & Palliative Care."
)

doc.add_heading("Rankings and Reputation", level=1)
doc.add_paragraph(
    "King's is a member of the Russel Group of research-intensive UK " \
    "universities and is consistently ranked among the top universities " \
    "globally, particularly strong in medicine, dentistry, law, the humanities, " \
    " and an active research department in computer science and artificial " \
    "intelligence. It was historically one of the founding colleges of the" \
    " federal University of London, though it has operated with significant " \
    "independence, including awarding its own degrees since 2007."
)

doc.add_heading("Student Life", level=1)
doc.add_paragraph(
    "King's College London Students' Union (KCLSU) runs student societies, "
    "sports clubs, and campus events across all five campuses. The university "
    "has a large international student population, reflecting London's global "
    "character, and offers halls of residence across multiple London locations."
)

doc.add_heading("Notable Alumni", level=1)
doc.add_paragraph(
    "King's College London counts among its alumni fourteen Nobel laureates across " \
    "physics, chemistry, meidcine and physiology, reflecting its historically " \
    "strong medical and scientific faculties. Physicist Preter Higgs, who proposed " \
    "the mechanism behind the Higgs boson, is among the most well known, alongside " \
    "Sir Michael Houghton (medicine, for reasearch into Hepatitis C) and Michael " \
    "Levitt (Chemistry, for work on computational modelling for chemical processes). " \
    "Archbishop Desmond Tutu, who studied theology at King's later recieved the " \
    "Nobel Peace Prize for his work against apartheid in South Africa."
)

doc.add_paragraph(
    "Beyond the sciences, King's alumni include prominent figures in politics, " \
    "law, literature, and the arts. Sir Keir Starmer studied at King's before " \
    "becoming a prominent figure in UK politics. In literature, alumni include " \
    "author and children's writer Michael Morpurgo and philosopher and writer " \
    "Alain de Botton. In music, Queen bassist John Deacon studied electronics " \
    "at King's. Olympic medalist Katherine Grainer (rowing) and Dina " \
    "Asher-smith (sprinting) also studied at the university."
)

doc.add_heading("Scale and Student Body", level=1)
doc.add_paragraph(
    "King's is one of the largest universities in the UK by enrolment, with a " \
    "total student population generally reported in the region of 40,000 or more, " \
    "including a substantial international makeup is frequently highlighted in " \
    "university rankings as a particular strength of King's global engagement, " \
    "alongside its extensive international research partnerships."
)

doc.add_heading("Rankings and Research Strength", level=1)
doc.add_paragraph(
    "King's consistently places among the top unversities globally across major " \
    "ranking systems such as QS, Times Higher Education, and the Academic Ranking " \
    "of World Universities, typically appearing within the global top 30-60 " \
    "Depending on the specific ranking and year, and generally among the top five " \
    "or six universities in the United Kingdom. It is particularly well regarded "
    "in subject-specific rankings for medicine, dentistry, nursing, law and " \
    "psychology, with a number of individual subjects ranking in the global top " \
    "15 in recent QS subject rankings."
)

doc.add_paragraph(
    "In the UK's Research Excellence Framework (REF), King's has ranked among the " \
    "top handful of UK universities for overall research power, reflecting both " \
    "the volume and quality of research produced across its faculties. Its annual " \
    "research income runs into the hundreds of millions of poinds, supporting " \
    "work across its medical, scientific, and social science departments, " \
    "including at its affiliated teaching hospitals."
)

doc.add_heading("Golden Triangle and London Context", level=1)
doc.add_paragraph(
    "King's is often described as part of the UK's 'Golden Triangle' of leading " \
    "research universities, a term referring to institutions concentrated in and " \
    "around London, Oxford, and Cambridge. Its location gives it close ties to " \
    "London's legal, financial medical, and cultural institutions, and London " \
    "itself is regularly ranked among the best cities in the world for students, " \
    "a factor King's frequently highlights in its own promotional material."
)

doc.add_heading("Sports and Athletics", level=1)
doc.add_paragraph(
    "King's fields a large number of student sports clubs under King's Sport and " \
    "the King's College London Student's Union (KCLSU), competing mainly through " \
    "British Universities and College Sport (BUCS), the main governing body "
    "for university sport in the UK. Clubs range from competitive teams playing in " \
    "structured BUCS leagues to purely recreational and social clubs aimed at " \
    "fitness and skill-building rather than ocmpetition."
)
doc.add_paragraph(
    "The centerpiece of King's sporting is the London Varsity Series, an " \
    "annual set of matches against University College London (UCL), King's " \
    "closest historical rival, help across multiple sports and drawing large " \
    "student crowds. A related fixture, the Macadam Cup, pits King's teams " \
    "against Guy's, King's and St Thomas' School of Medicine (GKT) teams."
)
doc.add_paragraph(
    "King's College London Rugby Football Club, founded in 1869, is one of the " \
    "university's oldest sports clubs and competes in BUCS leagues, playing its " \
    "home matches at New Malden Sports Ground. A separate, even older rugby " \
    "club associated with King's medical school, GKT Men's Rugby (dating to " \
    "1843), is also linked to the founding of the Rugby Football Union itself. " \
    "King's also has an active rowing club, King's College London Goat Club, " \
    "based on the Thames near Chiswick, with a history stretching back to the " \
    "1930s and past wins at the Henley Royal Regatta."
)

doc.add_heading("Computer Science, Informatics, and Artificial Intelliigence", level=1)
doc.add_paragraph(
    "Computer science and AI research and teaching at King's is centred in the " \
    "Department of Informatics, part of the Faculty of Natural, Mathematical & " \
    "Engineering Sciences. The department offers both undergraduate and " \
    "postgraduate taught degrees, alongside MPhil and PhD research degrees, and " \
    "maintains an active research profile with externally funded projects and " \
    "strong links to industry, government, and other academic institutions."
)
doc.add_paragraph(
    "King's has a dedicated King's Institute for Artificial Intelligence, which " \
    "coordinates AI-related research and events across the university, " \
    "including an anunual Festival of AI. Reserach within the department spans " \
    "areas such as machine learning, autonomous systems, robotics, and the " \
    "application of AI to fields including healthcare, finance, smart cities, " \
    "and energy systems, alongside a strong emphasis on AI sefety and " \
    "trustworthiness. Students on AI and computer science programmes have " \
    "opportunities to work on real industry projects, including partnerships " \
    "with major technologt companies."
)

doc.add_heading("Additional Notable Facts", level=1)
doc.add_paragraph(
    "King's has run large-scale philanthropic fundraising campaigns, including " \
    "one launched in 2010 that raised several hundred million pounds to " \
    "support scholarships, research, and a range of leadership and service " \
    "initiatives across the university. King's is also frequently cited " \
    "alongside Oxford, Cambridge, UCL, Imperial College London, and the London " \
    "School of Economics as one of the UK's most research-intensive " \
    "insitutions, forming part of what is often called the 'Golden Triangle' " \
    "of elite British universities concentrated around London, Oxford, and " \
    "Cambridge."
)

doc.save("data/kings_college_london.docx")
print("Created data/kings_college_london.docx")